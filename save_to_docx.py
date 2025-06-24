from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import datetime

def save_to_docx(text: str, filename: str = "Gereksinim_Dokumani.docx") -> None:
    doc = Document()

    # 📌 Üst Bilgi Tablosu
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'

    table.cell(0, 0).text = "Form Adı:"
    table.cell(0, 1).text = "Ürün Gereksinim Dokümanı"

    table.cell(1, 0).text = "Form Numarası:"
    table.cell(1, 1).text = "FRM-PRD-001"

    table.cell(2, 0).text = "Versiyon:"
    table.cell(2, 1).text = "1.0"

    table.cell(3, 0).text = "Hazırlayan:"
    table.cell(3, 1).text = "Otomatik AI Sistem"

    table.cell(4, 0).text = "Tarih:"
    table.cell(4, 1).text = datetime.today().strftime("%d.%m.%Y")

    doc.add_paragraph()  # boşluk bırak

    # 📄 Başlık
    title = doc.add_heading("Ürün Özellik Gereksinim Dokümanı", level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph()  # boşluk

    # 🧾 İçerik
    for line in text.splitlines():
        if line.strip() == "":
            doc.add_paragraph()
        elif line.strip().endswith(":") or line.strip().startswith("1.") or line.strip().startswith("###"):
            # Ana başlık gibi olan satırlar kalın ve büyük yazılsın
            para = doc.add_paragraph()
            run = para.add_run(line.strip())
            run.bold = True
            run.font.size = Pt(12)
        else:
            doc.add_paragraph(line.strip())

    # 💾 Kaydet
    doc.save(filename)
    print(f"✅ '{filename}' başarıyla kaydedildi.")
